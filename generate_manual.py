"""生成义乌小商品出海智能体-OPC 用户使用手册"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from datetime import datetime


def set_cell_shading(cell, color):
    """设置单元格底纹颜色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_table_with_style(doc, headers, rows, col_widths=None):
    """添加带样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(hdr_cells[i], "1F4E79")

    # 数据行
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = str(cell_text)
            for paragraph in row_cells[col_idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if row_idx % 2 == 0:
                set_cell_shading(row_cells[col_idx], "D6E4F0")

    # 设置列宽
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Cm(width)

    return table


def add_code_block(doc, code_text):
    """添加代码块（等宽字体）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    # 设置底纹
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
    pPr.append(shd)
    return p


def add_tip_box(doc, text, bold_prefix="💡 提示："):
    """添加提示框"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFF3CD"/>')
    pPr.append(shd)
    run_bold = p.add_run(bold_prefix)
    run_bold.bold = True
    run_bold.font.size = Pt(10)
    run_bold.font.color.rgb = RGBColor(0x85, 0x6A, 0x04)
    run_text = p.add_run(text)
    run_text.font.size = Pt(10)
    run_text.font.color.rgb = RGBColor(0x66, 0x50, 0x03)
    return p


def add_page_header(section, text):
    """添加页眉"""
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.text = text
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def add_page_footer(section):
    """添加页脚页码"""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    run2 = p.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instrText)
    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar2)


def generate_manual():
    doc = Document()

    # ============ 全局样式 ============
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 设置页面边距
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    add_page_header(section, "义乌小商品出海智能体-OPC 用户使用手册")
    add_page_footer(section)

    # ============ 封面 ============
    for _ in range(6):
        doc.add_paragraph()

    # 项目名称
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run("义乌小商品出海智能体-OPC")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run.font.name = "微软雅黑"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 英文名
    p_en = doc.add_paragraph()
    p_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_en = p_en.add_run("YiWu Global AI Agent")
    run_en.font.size = Pt(16)
    run_en.font.color.rgb = RGBColor(0x4A, 0x86, 0xC8)
    run_en.font.name = "Calibri"

    doc.add_paragraph()

    # 副标题
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("用户使用手册 V2.0 冠军版")
    run_sub.bold = True
    run_sub.font.size = Pt(20)
    run_sub.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    run_sub.font.name = "微软雅黑"
    run_sub.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    doc.add_paragraph()
    doc.add_paragraph()

    # 日期
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = p_date.add_run(f"2026年6月")
    run_date.font.size = Pt(14)
    run_date.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 参赛信息
    p_comp = doc.add_paragraph()
    p_comp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_comp = p_comp.add_run('2026"直通乌镇"全球互联网大赛 OPC特色赛')
    run_comp.font.size = Pt(12)
    run_comp.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 分页
    doc.add_page_break()

    # ============ 目录页 ============
    doc.add_heading("目  录", level=1)
    toc_items = [
        ("1", "产品简介"),
        ("2", "快速开始"),
        ("3", "首页仪表盘"),
        ("4", "市场洞察Agent"),
        ("5", "智能选品Agent"),
        ("6", "供应链匹配Agent"),
        ("7", "跨境内容生成Agent"),
        ("8", "合规助手Agent"),
        ("9", "智能客服Agent"),
        ("10", "政策复制Agent（V2新增重点）"),
        ("11", "全链路工作流"),
        ("12", "常见问题"),
        ("13", "技术支持"),
    ]
    for num, title in toc_items:
        p = doc.add_paragraph()
        run_num = p.add_run(f"  {num}.  ")
        run_num.bold = True
        run_num.font.size = Pt(12)
        run_title = p.add_run(title)
        run_title.font.size = Pt(12)

    doc.add_page_break()

    # ============ 第1章 产品简介 ============
    doc.add_heading("1  产品简介", level=1)

    doc.add_paragraph(
        "义乌小商品出海智能体-OPC是一款面向跨境电商的AI智能服务平台，依托义乌小商品城7.5万商户、"
        "210万+SKU的产业资源，集成7大AI Agent，为中小企业提供从市场分析到商品出海的全链路AI服务。"
    )

    doc.add_heading("1.1  7大AI Agent介绍", level=2)
    add_table_with_style(
        doc,
        ["Agent", "名称", "功能", "核心能力"],
        [
            ["🔍", "市场洞察", "全球市场趋势分析", "义乌指数解读、6大区域市场分析、品类趋势预测"],
            ["🎯", "智能选品", "多维度选品推荐", "竞争度评分、利润空间分析、义乌优势匹配"],
            ["🔗", "供应链匹配", "商铺与物流匹配", "7.5万商铺智能匹配、义新欧班列物流、1039模式对接"],
            ["✍️", "跨境内容生成", "多语言营销内容", "8语言4平台、SEO关键词、社媒文案、广告素材"],
            ["🛡️", "合规助手", "合规与关税查询", "5国合规要求、1039通关、RCEP优惠、关税计算"],
            ["💬", "智能客服", "多语言智能客服", "7×24小时、情绪识别、多轮对话、FAQ自动应答"],
            ["🏛️", "政策复制", "1039政策推广", "39城试点信息、政策解读、红利计算、成功案例"],
        ],
        col_widths=[2, 3, 4, 6],
    )

    doc.add_heading("1.2  义乌10大核心品类", level=2)
    add_table_with_style(
        doc,
        ["品类", "市场规模", "增长率", "义乌优势"],
        [
            ["日用百货", "580亿美元", "12.5%", "义乌一区、三区集中供应，品类齐全"],
            ["饰品配件", "320亿美元", "15.8%", "义乌五区饰品专区，全球最大饰品集散中心"],
            ["玩具", "950亿美元", "8.6%", "义乌一区玩具城，全球最大玩具批发市场之一"],
            ["文具办公用品", "280亿美元", "6.8%", "义乌三区文具专区，品类丰富"],
            ["针织品", "420亿美元", "10.2%", "义乌四区针织专区，全球最大袜子生产基地"],
            ["工艺品", "260亿美元", "9.5%", "义乌一区工艺品专区，圣诞用品占全球出口80%"],
            ["电子电器", "780亿美元", "14.3%", "义乌二区电子电器专区，小家电和LED灯饰品类齐全"],
            ["五金工具", "350亿美元", "7.5%", "义乌二区五金专区，工具品类齐全"],
            ["服装服饰", "1200亿美元", "11.2%", "义乌四区服装专区，快时尚供应链优势明显"],
            ["家居装饰", "450亿美元", "13.6%", "义乌五区家居装饰专区，新品更新速度快"],
        ],
        col_widths=[3, 3, 2, 7],
    )

    doc.add_heading("1.3  6大目标市场", level=2)
    add_table_with_style(
        doc,
        ["市场区域", "主要国家", "物流方式", "特点"],
        [
            ["欧洲", "德国、法国、西班牙、荷兰、波兰", "义新欧班列直达", "CE认证、REACH法规、消费力强"],
            ["中亚", "哈萨克斯坦、乌兹别克斯坦、吉尔吉斯斯坦", "义新欧班列7天直达", "EAC认证、EAEU关税优惠"],
            ["中东", "沙特、阿联酋、伊朗、土耳其", "海运+义新欧", "SABER认证、Halal认证、高消费"],
            ["东南亚", "印尼、泰国、越南、马来西亚", "海运", "RCEP优惠、SNI/TISI认证"],
            ["非洲", "埃及、尼日利亚、肯尼亚", "海运", "SONCAP认证、市场潜力大"],
            ["南美", "巴西、阿根廷、智利", "海运", "关税较高、市场增长快"],
        ],
        col_widths=[2.5, 4.5, 3.5, 5],
    )

    doc.add_page_break()

    # ============ 第2章 快速开始 ============
    doc.add_heading("2  快速开始", level=1)

    doc.add_heading("2.1  环境要求", level=2)
    add_table_with_style(
        doc,
        ["项目", "要求"],
        [
            ["Python", "3.12+"],
            ["Node.js", "18+"],
            ["npm", "9+"],
            ["操作系统", "Windows / macOS / Linux"],
            ["网络", "需要访问外部API"],
        ],
        col_widths=[4, 11],
    )

    doc.add_heading("2.2  后端启动", level=2)
    doc.add_paragraph("1. 进入后端目录：")
    add_code_block(doc, "cd demo")
    doc.add_paragraph("2. 安装Python依赖：")
    add_code_block(doc, "pip install -r requirements.txt")
    doc.add_paragraph("3. 配置环境变量：")
    add_code_block(doc, "cp .env.example .env\n# 编辑 .env 文件，填入必要配置")
    doc.add_paragraph("4. 启动后端服务：")
    add_code_block(doc, "uvicorn app.main:app --host 0.0.0.0 --port 8000 --reload")
    add_tip_box(doc, "后端服务启动后，可访问 http://localhost:8000/docs 查看API文档")

    doc.add_heading("2.3  前端启动", level=2)
    doc.add_paragraph("1. 进入前端目录：")
    add_code_block(doc, "cd demo/web")
    doc.add_paragraph("2. 安装前端依赖：")
    add_code_block(doc, "npm install")
    doc.add_paragraph("3. 启动开发服务器：")
    add_code_block(doc, "npm run dev")
    add_tip_box(doc, "前端启动后，访问 http://localhost:5173 即可使用")

    doc.add_heading("2.4  命令行模式", level=2)
    doc.add_paragraph("也可通过curl命令直接调用API：")
    add_code_block(doc, "# 获取市场洞察\ncurl \"http://localhost:8000/api/v1/market-insight?category=玩具&region=欧洲（义新欧班列直达）\"\n\n# 智能选品\ncurl \"http://localhost:8000/api/v1/smart-selection?category=饰品配件&budget=中&region=中东\"")

    doc.add_page_break()

    # ============ 第3章 首页仪表盘 ============
    doc.add_heading("3  首页仪表盘", level=1)

    doc.add_paragraph(
        "首页仪表盘是用户进入系统后的第一个页面，提供全局概览和快速导航功能。"
    )

    doc.add_heading("3.1  6个Agent状态", level=2)
    doc.add_paragraph(
        "首页顶部展示6个核心Agent的运行状态卡片，包括：市场洞察、智能选品、供应链匹配、"
        "跨境内容生成、合规助手、智能客服。每个卡片显示Agent名称、状态（在线/离线）和简要描述。"
    )
    add_tip_box(doc, "V2.0新增政策复制Agent，首页已同步展示其状态。")

    doc.add_heading("3.2  10大品类导航", level=2)
    doc.add_paragraph(
        "首页中部展示义乌10大核心品类导航卡片，点击任意品类可快速进入对应的市场洞察页面。"
        "品类包括：日用百货、饰品配件、玩具、文具办公用品、针织品、工艺品、电子电器、五金工具、服装服饰、家居装饰。"
    )

    doc.add_heading("3.3  义乌指数概览", level=2)
    doc.add_paragraph(
        "首页底部展示义乌指数概览，包括当前指数值（102.8）、变化趋势（上涨+1.35）和各品类指数分布。"
        "义乌指数是全球小商品价格的风向标，是选品决策的重要参考。"
    )

    doc.add_page_break()

    # ============ 第4章 市场洞察Agent ============
    doc.add_heading("4  市场洞察Agent", level=1)

    doc.add_paragraph(
        "市场洞察Agent基于义乌指数和多源数据，为用户提供全球市场趋势分析，帮助用户把握市场机会。"
    )

    doc.add_heading("4.1  使用步骤", level=2)
    steps = [
        "在首页或左侧导航栏点击「市场洞察」",
        "选择目标品类（如：玩具）",
        "选择目标市场区域（如：欧洲）",
        "点击「分析」按钮，等待AI生成报告",
        "查看报告详情，包括市场规模、增长率、热门产品等",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_heading("4.2  报告内容", level=2)
    add_table_with_style(
        doc,
        ["报告模块", "说明"],
        [
            ["市场规模", "目标品类在全球及各区域的市场规模"],
            ["增长率", "目标品类在各区域的市场增长率"],
            ["热门产品", "当前品类下的热销产品列表"],
            ["义乌优势", "义乌在该品类上的供应链优势"],
            ["义乌指数", "当前品类的义乌指数评分"],
            ["区域分布", "各目标市场的份额和增长情况"],
        ],
        col_widths=[4, 11],
    )

    doc.add_heading("4.3  数据来源", level=2)
    doc.add_paragraph("市场洞察Agent整合以下5大数据源：")
    for source in ["义乌小商品城 — 7.5万商户、210万+SKU实时数据", "义新欧班列 — 19条线路、50国160城物流数据",
                   "Amazon — 全球平台销售数据", "Alibaba.com — 国际站B2B数据", "行业报告 — 义乌指数及行业研究报告"]:
        doc.add_paragraph(source, style="List Bullet")

    add_tip_box(doc, "API端点：GET /api/v1/market-insight?category={品类}&region={区域}")

    doc.add_page_break()

    # ============ 第5章 智能选品Agent ============
    doc.add_heading("5  智能选品Agent", level=1)

    doc.add_paragraph(
        "智能选品Agent基于多维度数据分析，为用户推荐最优选品方案，降低选品风险，提高利润率。"
    )

    doc.add_heading("5.1  使用步骤", level=2)
    steps = [
        "点击左侧导航栏「智能选品」",
        "选择目标品类（如：饰品配件）",
        "选择预算等级（低/中/高）",
        "选择目标市场区域（如：中东）",
        "点击「推荐」按钮，获取AI选品建议",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_heading("5.2  分析内容", level=2)
    add_table_with_style(
        doc,
        ["分析维度", "说明"],
        [
            ["竞争度评分", "1-10分，分数越低竞争越小"],
            ["利润空间", "基于义乌源头价格与目标市场售价的利润分析"],
            ["义乌优势匹配", "该品类在义乌的供应链优势评分"],
            ["热门产品推荐", "基于市场数据推荐的热销产品"],
            ["价格区间", "义乌源头价格与目标市场建议售价"],
        ],
        col_widths=[4, 11],
    )

    doc.add_heading("5.3  评分说明", level=2)
    add_table_with_style(
        doc,
        ["评分等级", "分数范围", "建议"],
        [
            ["强烈推荐", "8-10分", "市场机会大，义乌优势明显，建议优先布局"],
            ["推荐", "6-8分", "市场机会较好，有一定竞争，建议谨慎进入"],
            ["一般", "4-6分", "市场竞争激烈，需差异化策略"],
            ["不推荐", "0-4分", "市场饱和或义乌优势不明显，不建议进入"],
        ],
        col_widths=[3, 3, 9],
    )

    add_tip_box(doc, "API端点：GET /api/v1/smart-selection?category={品类}&budget={预算}&region={区域}")

    doc.add_page_break()

    # ============ 第6章 供应链匹配Agent ============
    doc.add_heading("6  供应链匹配Agent", level=1)

    doc.add_paragraph(
        "供应链匹配Agent连接义乌7.5万商铺资源，为用户提供智能商铺匹配、采购对接、物流方案推荐服务。"
    )

    doc.add_heading("6.1  义乌7.5万商铺匹配", level=2)
    doc.add_paragraph(
        "系统根据用户选择的品类和需求，从义乌国际商贸城7.5万商铺中智能匹配最合适的供应商。"
        "匹配结果包括商铺名称、价格区间、起订量、所在区域等信息。"
    )
    add_table_with_style(
        doc,
        ["商贸城区", "主营品类", "商铺数量"],
        [
            ["一区", "玩具、工艺品", "12,000"],
            ["二区", "五金工具、电子电器", "15,000"],
            ["三区", "文具办公用品、日用百货", "13,000"],
            ["四区", "服装服饰、针织品", "14,000"],
            ["五区", "饰品配件、家居装饰", "11,000"],
        ],
        col_widths=[4, 6, 5],
    )

    doc.add_heading("6.2  采购对接", level=2)
    doc.add_paragraph(
        "匹配到合适的商铺后，用户可通过系统发起采购对接请求。系统将自动生成采购清单，"
        "包含产品名称、数量、价格、交期等信息，并支持在线沟通和订单跟踪。"
    )

    doc.add_heading("6.3  1039市场采购贸易模式", level=2)
    doc.add_paragraph(
        "1039市场采购贸易方式是义乌首创的贸易便利化政策，核心优势包括："
    )
    for advantage in [
        "增值税免征不退 — 无需取得增值税专用发票即可出口，合规成本降低80%",
        "简化申报 — 商品编码归并申报，报关项数从数十项缩减至5项以内",
        "通关便利化 — 优先查验、快速放行，出口通关时间压缩60%",
        "跨境人民币结算 — 规避汇率风险，简化外汇核销手续",
        "组柜拼箱 — 多商户共享集装箱，物流成本降低40%",
    ]:
        doc.add_paragraph(advantage, style="List Bullet")

    doc.add_heading("6.4  义新欧班列物流", level=2)
    doc.add_paragraph(
        "义新欧班列是义乌连接全球的铁路物流大通道，系统提供完整的班列信息查询服务。"
    )
    add_table_with_style(
        doc,
        ["线路", "运输天数", "班次频率", "20尺柜价格", "40尺柜价格"],
        [
            ["义乌-马德里", "21天", "每周3班", "$3,200", "$4,800"],
            ["义乌-伦敦", "18天", "每周2班", "$3,500", "$5,200"],
            ["义乌-德黑兰", "14天", "每周2班", "$2,800", "$4,200"],
            ["义乌-阿拉木图", "7天", "每周4班", "$1,800", "$2,800"],
            ["义乌-莫斯科", "10天", "每周3班", "$2,400", "$3,600"],
            ["义乌-明斯克", "12天", "每周2班", "$2,600", "$3,900"],
        ],
        col_widths=[3.5, 2.5, 2.5, 3, 3],
    )
    add_tip_box(doc, "义新欧班列比海运快2-3倍、比空运便宜60-80%，是小商品出海的最优物流方案。")

    add_tip_box(doc, "API端点：GET /api/v1/supply-chain/{category}?region={区域}&budget={预算}")

    doc.add_page_break()

    # ============ 第7章 跨境内容生成Agent ============
    doc.add_heading("7  跨境内容生成Agent", level=1)

    doc.add_paragraph(
        "跨境内容生成Agent支持8种语言、4大电商平台，一键生成专业的跨境电商营销内容。"
    )

    doc.add_heading("7.1  8语言4平台", level=2)
    add_table_with_style(
        doc,
        ["类别", "支持项"],
        [
            ["语言", "英语、德语、法语、西班牙语、阿拉伯语、俄语、哈萨克语、日语"],
            ["平台", "Amazon、Alibaba.com、TikTok Shop、Temu"],
        ],
        col_widths=[3, 12],
    )

    doc.add_heading("7.2  使用步骤", level=2)
    steps = [
        "点击左侧导航栏「内容生成」",
        "输入产品名称（如：益智积木）",
        "选择品类（如：玩具）",
        "选择目标平台（如：Amazon）",
        "选择目标语言（如：英语）",
        "点击「生成」按钮，获取AI生成的营销内容",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_heading("7.3  生成内容", level=2)
    add_table_with_style(
        doc,
        ["内容类型", "说明"],
        [
            ["产品标题", "SEO优化的多语言产品标题，包含核心关键词"],
            ["产品描述", "详细的产品描述，突出义乌源头优势"],
            ["卖点提炼", "4-6个核心卖点，配合图标展示"],
            ["SEO关键词", "按品类×语言生成的搜索关键词列表"],
            ["社媒文案", "适合社交媒体传播的短文案，含Hashtag"],
            ["广告素材", "广告标题、正文和CTA按钮文案"],
        ],
        col_widths=[4, 11],
    )

    add_tip_box(doc, "API端点：POST /api/v1/content/generate")

    doc.add_page_break()

    # ============ 第8章 合规助手Agent ============
    doc.add_heading("8  合规助手Agent", level=1)

    doc.add_paragraph(
        "合规助手Agent提供5大区域合规要求查询、1039通关指引、RCEP优惠查询和关税计算服务。"
    )

    doc.add_heading("8.1  5国合规要求", level=2)
    add_table_with_style(
        doc,
        ["国家", "核心认证", "关税范围", "VAT税率", "特殊要求"],
        [
            ["德国", "CE、RoHS、REACH", "3%-17%", "19%", "德语标签、VerpackG法规"],
            ["哈萨克斯坦", "EAC、GOST-K", "5%-15%", "12%", "俄语/哈萨克语标签"],
            ["沙特阿拉伯", "SABER、SASO", "5%-20%", "15%", "SABER系统注册、阿拉伯语标签"],
            ["印尼", "SNI、BPOM", "0%-40%", "11%", "印尼语标签、进口配额"],
            ["尼日利亚", "SONCAP、NAFDAC", "5%-35%", "7.5%", "SONCAP认证必须"],
        ],
        col_widths=[2.5, 3, 2, 2, 5.5],
    )

    doc.add_heading("8.2  1039模式通关", level=2)
    doc.add_paragraph("1039市场采购贸易通关流程：")
    steps = [
        "备案登记 — 在市场采购贸易综合管理系统中完成经营主体备案",
        "商品采购 — 在试点市场内采购商品，取得供货商户信息",
        "组货装箱 — 在指定监管场所完成组货装箱，生成装箱清单",
        "简化申报 — 通过综合管理系统进行简化申报，归并商品编码",
        "海关查验 — 海关实施便利化查验，优先放行",
        "出口通关 — 货物出口，完成通关手续",
        "收结汇 — 通过联网信息平台在线收结汇",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_heading("8.3  RCEP优惠", level=2)
    doc.add_paragraph("RCEP协定下中国出口至东盟国家的商品可享受以下优惠：")
    for benefit in [
        "90%以上税目产品最终零关税",
        "原产地累积规则降低享惠门槛",
        "经核准出口商自主声明原产地",
        "快件和快递货物6小时放行",
    ]:
        doc.add_paragraph(benefit, style="List Bullet")

    add_tip_box(doc, "API端点：GET /api/v1/compliance?category={品类}&target_country={国家}")
    add_tip_box(doc, "关税计算：POST /api/v1/tariff/calculate")

    doc.add_page_break()

    # ============ 第9章 智能客服Agent ============
    doc.add_heading("9  智能客服Agent", level=1)

    doc.add_paragraph(
        "智能客服Agent提供7×24小时多语言智能客服服务，支持情绪识别和FAQ自动应答。"
    )

    doc.add_heading("9.1  7×24多语言服务", level=2)
    doc.add_paragraph(
        "智能客服支持中文、英语、德语、法语、西班牙语、阿拉伯语、俄语、哈萨克语、日语等9种语言，"
        "覆盖所有目标市场区域。客服系统7×24小时在线，自动识别用户语言并切换对话模式。"
    )

    doc.add_heading("9.2  情绪识别", level=2)
    doc.add_paragraph(
        "系统内置情绪识别模块，可实时分析用户输入的情绪状态（积极/中性/消极/愤怒），"
        "当检测到用户情绪消极或愤怒时，自动升级为人工客服介入，确保用户体验。"
    )

    doc.add_heading("9.3  FAQ自动应答", level=2)
    doc.add_paragraph(
        "系统内置跨境电商常见问题知识库，覆盖以下主题："
    )
    for topic in [
        "1039市场采购贸易政策解读",
        "义新欧班列物流查询",
        "各品类合规认证要求",
        "RCEP关税优惠查询",
        "义乌国际商贸城采购指南",
        "跨境支付与结算",
        "售后服务与退换货",
    ]:
        doc.add_paragraph(topic, style="List Bullet")

    add_tip_box(doc, "API端点：POST /api/v1/customer-service/chat")
    add_tip_box(doc, "FAQ查询：GET /api/v1/customer-service/faq?category={品类}&language={语言}")

    doc.add_page_break()

    # ============ 第10章 政策复制Agent（V2新增重点） ============
    doc.add_heading("10  政策复制Agent（V2新增重点）", level=1)

    p_v2 = doc.add_paragraph()
    run_v2 = p_v2.add_run("⭐ 本章节为V2.0冠军版新增核心功能，是本次大赛的重点创新。")
    run_v2.bold = True
    run_v2.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    doc.add_paragraph(
        "政策复制Agent基于义乌1039市场采购贸易的成功经验，为全国39个试点城市提供政策解读、"
        "红利计算、成功案例分享和本地化推广方案，实现义乌发展经验的数字化复制推广。"
    )

    doc.add_heading("10.1  39城试点", level=2)
    doc.add_paragraph(
        "截至目前，国务院已在全国39个城市批准设立市场采购贸易方式试点。"
        "政策复制Agent提供完整的39城试点信息查询服务，包括各城市的主营品类、政策优惠和海关代码。"
    )
    add_table_with_style(
        doc,
        ["省份", "试点城市", "主营品类"],
        [
            ["浙江", "义乌、海宁、绍兴柯桥、湖州织里、台州路桥、温州瓯海、宁波江北、嘉兴平湖、杭州萧山、金华永康", "日用百货、饰品、纺织、童装、鞋类、小家电等"],
            ["福建", "泉州、厦门", "鞋服、石材、电子产品"],
            ["广东", "广州、深圳、佛山、东莞、中山、汕头", "服装、电子、陶瓷、家具、灯饰、玩具"],
            ["四川/重庆", "成都、重庆", "鞋类、家具、汽摩配件、电子产品"],
            ["云南/广西", "昆明、南宁", "花卉、珠宝、农产品、建材"],
            ["其他", "长沙、南昌、合肥、郑州、武汉、西安、兰州、乌鲁木齐等", "工程机械、纺织品、家电、农产品等"],
        ],
        col_widths=[2.5, 6.5, 6],
    )
    add_tip_box(doc, "API端点：GET /api/v1/policy-replication/cities")

    doc.add_heading("10.2  政策解读", level=2)
    doc.add_paragraph("1039市场采购贸易政策核心要点：")
    add_table_with_style(
        doc,
        ["政策要点", "说明", "受益程度"],
        [
            ["增值税免征不退", "出口货物免征增值税，无需取得增值税专用发票", "高"],
            ["简化申报", "商品编码归并申报，报关项数缩减至5项以内", "高"],
            ["通关便利化", "优先查验、快速放行，通关时间压缩60%", "高"],
            ["跨境人民币结算", "规避汇率风险，简化外汇核销手续", "中"],
            ["组柜拼箱", "多商户共享集装箱，物流成本降低40%", "中"],
            ["在线收结汇", "联网信息平台在线收结汇，资金到账快", "中"],
        ],
        col_widths=[4, 8, 3],
    )
    add_tip_box(doc, "API端点：GET /api/v1/policy-replication/policy-guide")

    doc.add_heading("10.3  红利计算", level=2)
    doc.add_paragraph(
        "政策复制Agent内置红利计算器，用户输入年出口额、品类和所在城市，"
        "即可自动计算1039模式下的政策红利，包括："
    )
    for item in [
        "增值税节省金额 — 对比一般贸易需取得增值税专用发票的成本",
        "合规成本节省 — 1039模式合规成本率仅0.5%，一般贸易为3%",
        "通关时间节省 — 1039通关1天 vs 一般贸易3天",
        "物流成本节省 — 组柜拼箱可节省30%物流成本",
        "所得税优惠 — 核定征收应税所得率统一按5%核定",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_tip_box(doc, "API端点：POST /api/v1/policy-replication/calculate-benefit")

    doc.add_heading("10.4  成功案例", level=2)
    doc.add_paragraph("系统内置5个义乌成功案例，供其他城市参考借鉴：")
    add_table_with_style(
        doc,
        ["案例", "品类", "目标市场", "年出口额", "核心策略"],
        [
            ["饰品出海中东", "饰品配件", "阿联酋、沙特", "1200万美元", "1039简化申报+义新欧班列"],
            ["玩具出口东南亚", "玩具", "印尼、泰国、越南", "800万美元", "RCEP+1039双政策红利"],
            ["日用百货出口欧洲", "日用百货", "德国、法国、西班牙", "2500万美元", "义新欧班列+1039模式"],
            ["针织品出口非洲", "针织品", "尼日利亚、肯尼亚", "600万美元", "小批量多品种1039典范"],
            ["电子电器出口中亚", "电子电器", "哈萨克斯坦等", "1500万美元", "EAEU+1039政策组合拳"],
        ],
        col_widths=[3, 2.5, 3.5, 2.5, 4],
    )
    add_tip_box(doc, "API端点：GET /api/v1/policy-replication/cases")

    doc.add_page_break()

    # ============ 第11章 全链路工作流 ============
    doc.add_heading("11  全链路工作流", level=1)

    doc.add_paragraph(
        "全链路工作流将7大AI Agent串联，实现从市场分析到商品出海的一站式服务。"
    )

    doc.add_heading("11.1  7步串联", level=2)
    add_table_with_style(
        doc,
        ["步骤", "Agent", "功能"],
        [
            ["1", "市场洞察", "分析目标品类和区域的市场趋势"],
            ["2", "智能选品", "基于市场分析推荐最优选品方案"],
            ["3", "供应链匹配", "匹配义乌商铺和物流方案"],
            ["4", "跨境内容生成", "生成多语言营销内容"],
            ["5", "合规查询", "查询目标市场合规要求和关税"],
            ["6", "智能客服", "提供售前售后多语言客服"],
            ["7", "政策复制", "计算1039政策红利，提供推广方案"],
        ],
        col_widths=[2, 4, 9],
    )

    doc.add_heading("11.2  API调用", level=2)
    doc.add_paragraph("全链路工作流通过单一API调用即可完成：")
    add_code_block(doc,
        'curl -X POST "http://localhost:8000/api/v1/pipeline" \\\n'
        '  -H "Content-Type: application/json" \\\n'
        '  -d \'{\n'
        '    "category": "玩具",\n'
        '    "region": "欧洲（义新欧班列直达）",\n'
        '    "budget": "中",\n'
        '    "target_country": "德国",\n'
        '    "platform": "amazon",\n'
        '    "target_language": "de"\n'
        '  }\''
    )

    add_tip_box(doc, "全链路工作流会依次调用7个Agent，返回完整的出海方案报告。")

    doc.add_page_break()

    # ============ 第12章 常见问题 ============
    doc.add_heading("12  常见问题", level=1)

    faqs = [
        ("Q1：什么是1039市场采购贸易方式？",
         "1039是海关监管代码，指在经认定的市场集聚区采购商品，单票报关单商品货值15万美元（含）以下，"
         "可直接办理出口通关。核心优势是增值税免征、简化申报、通关便利化。"),
        ("Q2：义新欧班列的运输时效如何？",
         "义新欧班列覆盖19条线路、50国160城。到中亚7天、到俄罗斯10天、到欧洲14-21天。"
         "比海运快2-3倍、比空运便宜60-80%。"),
        ("Q3：哪些品类适合通过义乌出海？",
         "义乌10大核心品类均适合出海：日用百货、饰品配件、玩具、文具办公用品、针织品、工艺品、"
         "电子电器、五金工具、服装服饰、家居装饰。其中玩具、饰品、工艺品是义乌的传统优势品类。"),
        ("Q4：如何获取义乌指数数据？",
         "通过市场洞察Agent或直接调用API（GET /api/v1/yiwu-index）即可获取义乌指数数据，"
         "包括当前指数值、变化趋势和各品类指数分布。"),
        ("Q5：1039模式需要什么条件？",
         "主要条件：①在经认定的市场采购贸易试点区域内采购；②经由海关监管的采购地出口；"
         "③单票报关单商品货值15万美元（含）以下；④在市场采购贸易综合管理系统中备案。"),
        ("Q6：跨境内容生成支持哪些语言？",
         "支持8种语言：英语、德语、法语、西班牙语、阿拉伯语、俄语、哈萨克语、日语。"
         "覆盖欧洲、中亚、中东、东南亚等主要目标市场。"),
        ("Q7：如何计算1039政策红利？",
         "通过政策复制Agent的红利计算器，输入年出口额、品类和所在城市，系统将自动计算"
         "增值税节省、合规成本节省、通关时间节省和物流成本节省。"),
        ("Q8：39城试点城市有哪些？",
         "包括浙江10城（义乌、海宁等）、福建2城、广东6城、四川/重庆各1城、云南/广西各1城，"
         "以及长沙、南昌、合肥、郑州、武汉、西安等中西部城市，共39个。"),
        ("Q9：智能客服支持哪些语言？",
         "智能客服支持9种语言：中文、英语、德语、法语、西班牙语、阿拉伯语、俄语、哈萨克语、日语。"
         "系统自动识别用户语言并切换对话模式。"),
        ("Q10：全链路工作流如何使用？",
         "通过API（POST /api/v1/pipeline）一次性提交品类、区域、预算、目标国家、平台和语言参数，"
         "系统将自动串联7大Agent，返回完整的出海方案报告。"),
    ]

    for question, answer in faqs:
        p_q = doc.add_paragraph()
        run_q = p_q.add_run(question)
        run_q.bold = True
        run_q.font.size = Pt(11)
        run_q.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        p_a = doc.add_paragraph(answer)
        p_a.paragraph_format.left_indent = Cm(0.5)
        p_a.paragraph_format.space_after = Pt(8)

    doc.add_page_break()

    # ============ 第13章 技术支持 ============
    doc.add_heading("13  技术支持", level=1)

    doc.add_heading("13.1  联系方式", level=2)
    add_table_with_style(
        doc,
        ["渠道", "信息"],
        [
            ["邮箱", "contact@yiwu-global-ai.com"],
            ["GitHub", "https://github.com"],
            ["参赛团队", "义乌小商品出海智能体团队"],
            ["参赛赛道", '2026"直通乌镇"全球互联网大赛 OPC特色赛'],
        ],
        col_widths=[4, 11],
    )

    doc.add_heading("13.2  API文档", level=2)
    doc.add_paragraph("启动后端服务后，访问以下地址查看完整API文档：")
    add_code_block(doc, "http://localhost:8000/docs")

    doc.add_heading("13.3  系统状态", level=2)
    doc.add_paragraph("可通过以下API查看系统运行状态：")
    add_code_block(doc, "curl http://localhost:8000/api/v1/status")

    doc.add_heading("13.4  版本历史", level=2)
    add_table_with_style(
        doc,
        ["版本", "日期", "更新内容"],
        [
            ["V1.0", "2026年3月", "初始版本，6大Agent核心功能"],
            ["V2.0", "2026年6月", "冠军版：新增政策复制Agent、39城试点、红利计算、成功案例"],
        ],
        col_widths=[3, 4, 8],
    )

    # 保存文档
    output_path = r"D:\YiWuInternetCompetition\用户使用手册.docx"
    doc.save(output_path)
    print(f"用户使用手册已生成：{output_path}")


if __name__ == "__main__":
    generate_manual()
