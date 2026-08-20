#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将 商业计划书.md 的内容生成为专业排版的Word文档
"""

import re
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from pathlib import Path

# ── 常量 ──────────────────────────────────────────────
YIWU_RED = RGBColor(0xD4, 0x27, 0x2C)       # 义乌红 #D4272C
FONT_BODY = "微软雅黑"
FONT_HEADING = "黑体"
MD_PATH = Path(__file__).parent / "商业计划书.md"
DOCX_PATH = Path(__file__).parent / "商业计划书.docx"


# ── 辅助函数 ──────────────────────────────────────────

def set_cell_shading(cell, color_hex: str):
    """设置单元格底纹颜色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_table_borders(table):
    """为表格设置边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def add_formatted_text(paragraph, text: str, base_font_name=FONT_BODY, base_font_size=Pt(10.5), base_color=None):
    """向段落添加带格式的文本，支持 **加粗** 和 `代码` 标记"""
    # 拆分文本中的加粗和代码标记
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = base_font_name
            run.font.size = base_font_size
            run._element.rPr.rFonts.set(qn('w:eastAsia'), base_font_name)
            if base_color:
                run.font.color.rgb = base_color
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            run = paragraph.add_run(part)
            run.font.name = base_font_name
            run.font.size = base_font_size
            if base_color:
                run.font.color.rgb = base_color
            run._element.rPr.rFonts.set(qn('w:eastAsia'), base_font_name)


def add_heading(doc, text: str, level: int):
    """添加标题并设置字体"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = FONT_HEADING
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEADING)
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = YIWU_RED
        elif level == 2:
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif level == 3:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return heading


def add_paragraph_with_style(doc, text: str, font_size=Pt(10.5), bold=False, alignment=None, space_after=Pt(6), color=None):
    """添加正文段落"""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = Pt(20)
    add_formatted_text(p, text, base_font_size=font_size, base_color=color)
    if bold:
        for run in p.runs:
            run.bold = True
    return p


def add_bullet_item(doc, text: str, level=0):
    """添加列表项"""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.8)
    add_formatted_text(p, text, base_font_size=Pt(10.5))
    return p


def add_code_block(doc, lines: list):
    """添加代码块"""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = Pt(16)
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        # 设置段落底纹
        pPr = p._element.get_or_add_pPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
        pPr.append(shading)


def add_table(doc, headers: list, rows: list):
    """添加带格式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header.strip())
        run.bold = True
        run.font.name = FONT_BODY
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
        set_cell_shading(cell, "D4272C")

    # 数据行
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            text = cell_text.strip()
            add_formatted_text(p, text, base_font_size=Pt(9.5))
            # 交替行底纹
            if r_idx % 2 == 1:
                set_cell_shading(cell, "FFF3F3")

    # 设置列宽自适应
    table.autofit = True
    doc.add_paragraph()  # 表格后空行
    return table


def add_page_header(section, text: str):
    """添加页眉"""
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
    # 页眉下方加线
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


def add_page_footer(section):
    """添加页脚（页码）"""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 页码上方加线
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)
    run = p.add_run("第 ")
    run.font.name = FONT_BODY
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
    # 插入页码字段
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run1 = p.add_run()
    run1._element.append(fldChar1)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2 = p.add_run()
    run2._element.append(instrText)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3 = p.add_run()
    run3._element.append(fldChar2)
    run4 = p.add_run(" 页")
    run4.font.name = FONT_BODY
    run4.font.size = Pt(8)
    run4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run4._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)


# ── 封面页 ──────────────────────────────────────────

def create_cover_page(doc):
    """创建封面页"""
    # 多个空行将内容推到页面中下部
    for _ in range(6):
        doc.add_paragraph()

    # 项目名称
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("义乌小商品出海智能体-OPC")
    run.bold = True
    run.font.name = FONT_HEADING
    run.font.size = Pt(32)
    run.font.color.rgb = YIWU_RED
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEADING)

    # 副标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("商业计划书 V2.0 冠军版")
    run.bold = True
    run.font.name = FONT_HEADING
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEADING)

    # 分隔线
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run("━" * 30)
    run.font.color.rgb = YIWU_RED
    run.font.size = Pt(12)

    # 参赛信息
    info_lines = [
        "参赛赛道：2026\"直通乌镇\"全球互联网大赛 OPC特色赛",
        "创始人：冯亦根（浙江省首批产业教授、高级工程师、工信部认证首席数据官）",
        "版本：V2.0 冠军版",
        "日期：2026年6月",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(line)
        run.font.name = FONT_BODY
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)

    # 分页
    doc.add_page_break()


# ── 目录页 ──────────────────────────────────────────

def create_toc_page(doc):
    """创建目录页"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run("目  录")
    run.bold = True
    run.font.name = FONT_HEADING
    run.font.size = Pt(22)
    run.font.color.rgb = YIWU_RED
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEADING)

    toc_items = [
        ("一、项目概述", 1),
        ("1.1 国家战略背景：义乌发展经验", 2),
        ("1.2 项目定位", 2),
        ("1.3 核心价值主张", 2),
        ("1.4 三级增长飞轮", 2),
        ("1.5 OPC模式说明", 2),
        ("1.6 项目愿景", 2),
        ("二、市场分析", 1),
        ("2.1 全球小商品市场规模", 2),
        ("2.2 中国小商品出口数据", 2),
        ("2.3 义乌市场核心数据", 2),
        ("2.4 义新欧班列数据", 2),
        ("2.5 目标市场分析", 2),
        ("2.6 义乌综试区政策优势", 2),
        ("2.7 1039模式全国复制推广", 2),
        ('2.8 义乌发展经验\u201c六个坚持\u201d', 2),
        ("三、痛点分析", 1),
        ("3.1 义乌商户出海五大痛点", 2),
        ("3.2 传统解决方案的局限性", 2),
        ("四、产品方案", 1),
        ("4.1 七大AI Agent详细介绍", 2),
        ("4.2 全链路工作流：7个Agent自动串联", 2),
        ("4.3 义乌独有壁垒", 2),
        ("五、技术架构", 1),
        ("5.1 整体架构", 2),
        ("5.2 LangGraph多Agent协作框架", 2),
        ("5.3 FastAPI异步后端", 2),
        ("5.4 React 18前端", 2),
        ("5.5 数据源", 2),
        ("5.6 技术创新点", 2),
        ("六、商业模式", 1),
        ("6.1 产品定价", 2),
        ("6.2 收入结构", 2),
        ("6.3 收入预测（12个月）", 2),
        ("6.4 单位经济模型", 2),
        ("6.5 39城复制推广收入模型", 2),
        ("七、竞争优势", 1),
        ("7.1 竞品对比分析", 2),
        ("7.2 竞争优势详解", 2),
        ("7.3 护城河建设", 2),
        ("八、团队与OPC模式", 1),
        ("8.1 创始人：冯亦根", 2),
        ("8.2 OPC模式深度解读", 2),
        ("8.3 成本对比：OPC模式 vs 传统团队", 2),
        ("九、融资计划", 1),
        ("9.1 融资需求", 2),
        ("9.2 资金使用计划", 2),
        ("9.3 盈亏平衡预测", 2),
        ("十、发展规划", 1),
        ("10.1 第一阶段：MVP验证", 2),
        ("10.2 第二阶段：增长验证", 2),
        ("10.3 第三阶段：规模化", 2),
        ("10.4 长期愿景（1-3年）", 2),
        ("附录", 1),
    ]

    for title, level in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = Pt(22)
        if level == 1:
            p.paragraph_format.left_indent = Cm(0)
            run = p.add_run(title)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = YIWU_RED
        else:
            p.paragraph_format.left_indent = Cm(1.2)
            run = p.add_run(title)
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        run.font.name = FONT_BODY
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)

    doc.add_page_break()


# ── Markdown 解析与正文生成 ──────────────────────────

def parse_and_render(doc, md_text: str):
    """解析Markdown文本并渲染到Word文档"""
    lines = md_text.split('\n')
    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 代码块开始/结束
        if stripped.startswith('```'):
            if in_code_block:
                # 代码块结束
                add_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # 空行
        if not stripped:
            i += 1
            continue

        # 水平线
        if stripped == '---' or stripped == '***' or stripped == '___':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("─" * 50)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)
            i += 1
            continue

        # 标题
        heading_match = re.match(r'^(#{1,3})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            title_text = heading_match.group(2).strip()
            add_heading(doc, title_text, level)
            i += 1
            continue

        # 引用块
        if stripped.startswith('>'):
            quote_text = stripped.lstrip('> ').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = Pt(18)
            # 引用左边框
            pPr = p._element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                '  <w:left w:val="single" w:sz="12" w:space="4" w:color="D4272C"/>'
                '</w:pBdr>'
            )
            pPr.append(pBdr)
            add_formatted_text(p, quote_text, base_font_size=Pt(10), base_color=RGBColor(0x55, 0x55, 0x55))
            i += 1
            continue

        # 表格
        if '|' in stripped and i + 1 < len(lines) and re.match(r'^\s*\|[\s\-:|]+\|\s*$', lines[i + 1].strip()):
            # 解析表格
            header_line = stripped
            # 跳过分隔行
            i += 2
            data_lines = []
            while i < len(lines) and '|' in lines[i].strip() and lines[i].strip():
                data_lines.append(lines[i].strip())
                i += 1

            # 解析表头和数据
            headers = [c.strip() for c in header_line.split('|') if c.strip()]
            rows = []
            for dl in data_lines:
                cells = [c.strip() for c in dl.split('|') if c.strip() or True]
                # 保留空单元格
                parts = dl.split('|')
                cells = [p.strip() for p in parts]
                # 去掉首尾空元素
                if cells and cells[0] == '':
                    cells = cells[1:]
                if cells and cells[-1] == '':
                    cells = cells[:-1]
                if len(cells) == len(headers):
                    rows.append(cells)

            add_table(doc, headers, rows)
            continue

        # 列表项
        list_match = re.match(r'^(\s*)[-*]\s+(.+)$', line)
        if list_match:
            indent = len(list_match.group(1))
            level = indent // 2
            text = list_match.group(2)
            add_bullet_item(doc, text, level)
            i += 1
            continue

        # 有序列表
        olist_match = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
        if olist_match:
            text = olist_match.group(2)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = Pt(18)
            add_formatted_text(p, text, base_font_size=Pt(10.5))
            i += 1
            continue

        # 普通段落
        # 合并连续的普通文本行
        para_text = stripped
        while i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            if (not next_line or
                next_line.startswith('#') or
                next_line.startswith('>') or
                next_line.startswith('- ') or
                next_line.startswith('* ') or
                next_line.startswith('```') or
                next_line == '---' or
                '|' in next_line or
                re.match(r'^\d+\.\s+', next_line)):
                break
            i += 1
            para_text += ' ' + next_line

        add_paragraph_with_style(doc, para_text)
        i += 1


# ── 主函数 ──────────────────────────────────────────

def main():
    # 读取Markdown文件
    md_text = MD_PATH.read_text(encoding='utf-8')

    # 去掉文件开头的标题行（封面页单独处理）
    # 找到第一个 --- 分隔符，去掉封面信息部分
    parts = md_text.split('---', 1)
    if len(parts) > 1:
        body_text = parts[1]
    else:
        body_text = md_text

    # 创建文档
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_BODY
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)

    # 设置页面边距
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    # 添加页眉页脚
    add_page_header(section, "义乌小商品出海智能体-OPC 商业计划书")
    add_page_footer(section)

    # 封面页
    create_cover_page(doc)

    # 目录页
    create_toc_page(doc)

    # 解析并渲染正文
    parse_and_render(doc, body_text)

    # 保存文档
    doc.save(str(DOCX_PATH))
    print(f"✅ Word文档已成功生成：{DOCX_PATH}")
    print(f"   文件大小：{DOCX_PATH.stat().st_size / 1024:.1f} KB")


if __name__ == '__main__':
    main()
