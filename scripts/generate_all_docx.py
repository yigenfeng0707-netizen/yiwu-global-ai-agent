#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
批量将 Markdown 文档生成为专业排版的 Word 文档
文档列表：竞品分析、技术文档、财务预测与商业计划、里程碑规划、风险应对预案、路演演讲稿、答辩QA手册、视频脚本、部署指南
"""

import re
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from pathlib import Path

# ── 常量 ──────────────────────────────────────────────
YIWU_RED = RGBColor(0xD4, 0x27, 0x2C)       # 义乌红 #D4272C
FONT_BODY = "微软雅黑"
FONT_HEADING = "黑体"
BASE_DIR = Path(__file__).parent

# 需要转换的文档列表
DOC_LIST = [
    "竞品分析",
    "技术文档",
    "财务预测与商业计划",
    "里程碑规划",
    "风险应对预案",
    "路演演讲稿",
    "答辩QA手册",
    "视频脚本",
    "部署指南",
]


# ── 辅助函数 ──────────────────────────────────────────

def set_cell_shading(cell, color_hex: str):
    """设置单元格底纹颜色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_table_borders(table):
    """为表格设置边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def add_formatted_text(paragraph, text: str, base_font_name=FONT_BODY, base_font_size=Pt(10.5), base_color=None):
    """向段落添加带格式的文本，支持 **加粗** 和 `代码` 标记"""
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = base_font_name
            run.font.size = base_font_size
            run._element.rPr.rFonts.set(qn('w:eastAsia'), base_font_name)
            if base_color:
                run.font.color.rgb = base_color
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            run = paragraph.add_run(part)
            run.font.name = base_font_name
            run.font.size = base_font_size
            if base_color:
                run.font.color.rgb = base_color
            run._element.rPr.rFonts.set(qn('w:eastAsia'), base_font_name)


def add_heading(doc, text: str, level: int):
    """添加标题并设置字体"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = FONT_HEADING
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEADING)
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = YIWU_RED
        elif level == 2:
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif level == 3:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return heading


def add_paragraph_with_style(doc, text: str, font_size=Pt(10.5), bold=False, alignment=None, space_after=Pt(6), color=None):
    """添加正文段落"""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = Pt(20)
    add_formatted_text(p, text, base_font_size=font_size, base_color=color)
    if bold:
        for run in p.runs:
            run.bold = True
    return p


def add_bullet_item(doc, text: str, level=0):
    """添加列表项"""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.8)
    add_formatted_text(p, text, base_font_size=Pt(10.5))
    return p


def add_code_block(doc, lines: list):
    """添加代码块"""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = Pt(16)
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        pPr = p._element.get_or_add_pPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
        pPr.append(shading)


def add_table(doc, headers: list, rows: list):
    """添加带格式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header.strip())
        run.bold = True
        run.font.name = FONT_BODY
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
        set_cell_shading(cell, "D4272C")

    # 数据行
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            text = cell_text.strip()
            add_formatted_text(p, text, base_font_size=Pt(9.5))
            if r_idx % 2 == 1:
                set_cell_shading(cell, "FFF3F3")

    table.autofit = True
    doc.add_paragraph()
    return table


def add_page_header(section, text: str):
    """添加页眉"""
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


def add_page_footer(section):
    """添加页脚（页码）"""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)
    run = p.add_run("第 ")
    run.font.name = FONT_BODY
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run1 = p.add_run()
    run1._element.append(fldChar1)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2 = p.add_run()
    run2._element.append(instrText)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3 = p.add_run()
    run3._element.append(fldChar2)
    run4 = p.add_run(" 页")
    run4.font.name = FONT_BODY
    run4.font.size = Pt(8)
    run4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run4._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)


# ── 封面页 ──────────────────────────────────────────

def create_cover_page(doc, doc_name: str):
    """创建封面页，标题为文档名称"""
    for _ in range(6):
        doc.add_paragraph()

    # 项目名称
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("义乌小商品出海智能体-OPC")
    run.bold = True
    run.font.name = FONT_HEADING
    run.font.size = Pt(32)
    run.font.color.rgb = YIWU_RED
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEADING)

    # 文档名称副标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(doc_name)
    run.bold = True
    run.font.name = FONT_HEADING
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEADING)

    # 分隔线
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run("━" * 30)
    run.font.color.rgb = YIWU_RED
    run.font.size = Pt(12)

    # 参赛信息
    info_lines = [
        "参赛赛道：2026\"直通乌镇\"全球互联网大赛 OPC特色赛",
        "创始人：冯亦根（浙江省首批产业教授、高级工程师、工信部认证首席数据官）",
        "日期：2026年6月",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(line)
        run.font.name = FONT_BODY
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)

    doc.add_page_break()


# ── Markdown 解析与正文生成 ──────────────────────────

def parse_and_render(doc, md_text: str):
    """解析Markdown文本并渲染到Word文档"""
    lines = md_text.split('\n')
    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 代码块开始/结束
        if stripped.startswith('```'):
            if in_code_block:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # 空行
        if not stripped:
            i += 1
            continue

        # 水平线
        if stripped == '---' or stripped == '***' or stripped == '___':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("─" * 50)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)
            i += 1
            continue

        # 标题
        heading_match = re.match(r'^(#{1,3})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            title_text = heading_match.group(2).strip()
            add_heading(doc, title_text, level)
            i += 1
            continue

        # 引用块
        if stripped.startswith('>'):
            quote_text = stripped.lstrip('> ').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = Pt(18)
            pPr = p._element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                '  <w:left w:val="single" w:sz="12" w:space="4" w:color="D4272C"/>'
                '</w:pBdr>'
            )
            pPr.append(pBdr)
            add_formatted_text(p, quote_text, base_font_size=Pt(10), base_color=RGBColor(0x55, 0x55, 0x55))
            i += 1
            continue

        # 表格
        if '|' in stripped and i + 1 < len(lines) and re.match(r'^\s*\|[\s\-:|]+\|\s*$', lines[i + 1].strip()):
            header_line = stripped
            i += 2
            data_lines = []
            while i < len(lines) and '|' in lines[i].strip() and lines[i].strip():
                data_lines.append(lines[i].strip())
                i += 1

            headers = [c.strip() for c in header_line.split('|') if c.strip()]
            rows = []
            for dl in data_lines:
                parts = dl.split('|')
                cells = [p.strip() for p in parts]
                if cells and cells[0] == '':
                    cells = cells[1:]
                if cells and cells[-1] == '':
                    cells = cells[:-1]
                if len(cells) == len(headers):
                    rows.append(cells)

            add_table(doc, headers, rows)
            continue

        # 列表项
        list_match = re.match(r'^(\s*)[-*]\s+(.+)$', line)
        if list_match:
            indent = len(list_match.group(1))
            level = indent // 2
            text = list_match.group(2)
            add_bullet_item(doc, text, level)
            i += 1
            continue

        # 有序列表
        olist_match = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
        if olist_match:
            text = olist_match.group(2)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = Pt(18)
            add_formatted_text(p, text, base_font_size=Pt(10.5))
            i += 1
            continue

        # 普通段落
        para_text = stripped
        while i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            if (not next_line or
                next_line.startswith('#') or
                next_line.startswith('>') or
                next_line.startswith('- ') or
                next_line.startswith('* ') or
                next_line.startswith('```') or
                next_line == '---' or
                '|' in next_line or
                re.match(r'^\d+\.\s+', next_line)):
                break
            i += 1
            para_text += ' ' + next_line

        add_paragraph_with_style(doc, para_text)
        i += 1


# ── 单文档生成 ──────────────────────────────────────

def generate_docx(doc_name: str):
    """将单个Markdown文档生成为Word文档"""
    md_path = BASE_DIR / f"{doc_name}.md"
    docx_path = BASE_DIR / f"{doc_name}.docx"

    if not md_path.exists():
        print(f"⚠️  跳过 {doc_name}：Markdown文件不存在 ({md_path})")
        return False

    md_text = md_path.read_text(encoding='utf-8')

    # 去掉文件开头的标题行（封面页单独处理）
    parts = md_text.split('---', 1)
    if len(parts) > 1:
        body_text = parts[1]
    else:
        body_text = md_text

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_BODY
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)

    # 设置页面边距
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    # 添加页眉页脚
    add_page_header(section, f"义乌小商品出海智能体-OPC {doc_name}")
    add_page_footer(section)

    # 封面页
    create_cover_page(doc, doc_name)

    # 解析并渲染正文
    parse_and_render(doc, body_text)

    # 保存文档
    doc.save(str(docx_path))
    size_kb = docx_path.stat().st_size / 1024
    print(f"✅ {doc_name}.docx 已生成 ({size_kb:.1f} KB)")
    return True


# ── 主函数 ──────────────────────────────────────────

def main():
    print("=" * 60)
    print("  批量生成 Word 文档")
    print("  义乌小商品出海智能体-OPC")
    print("=" * 60)
    print()

    success_count = 0
    fail_count = 0

    for doc_name in DOC_LIST:
        try:
            if generate_docx(doc_name):
                success_count += 1
            else:
                fail_count += 1
        except Exception as e:
            print(f"❌ {doc_name}.docx 生成失败：{e}")
            fail_count += 1

    print()
    print("=" * 60)
    print(f"  完成！成功 {success_count} 个，失败 {fail_count} 个")
    print("=" * 60)


if __name__ == '__main__':
    main()
