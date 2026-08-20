#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成OPC赛道商业计划书 - 专门针对OPC赛道要求
重点：AI提效可视化对比、AI提效深度与广度、产品矩阵
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from pathlib import Path

# ── 常量 ──────────────────────────────────────────────
YIWU_RED = RGBColor(0xD4, 0x27, 0x2C)  # 义乌红 #D4272C
YIWU_RED_HEX = "D4272C"
LIGHT_RED_HEX = "FFF3F3"
FONT_BODY = "微软雅黑"
FONT_HEADING = "黑体"
DOCX_PATH = Path(__file__).parent / "OPC赛道商业计划书.docx"


# ── 辅助函数 ──────────────────────────────────────────

def set_cell_shading(cell, color_hex: str):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def add_run(paragraph, text, font_name=FONT_BODY, font_size=Pt(10.5), bold=False, color=None):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    if color:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return run


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = FONT_HEADING
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEADING)
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0xD4, 0x27, 0x2C)
        elif level == 2:
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif level == 3:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return heading


def add_para(doc, text, font_size=Pt(10.5), bold=False, alignment=None, space_after=Pt(6), color=None):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = Pt(20)
    add_run(p, text, font_size=font_size, bold=bold, color=color)
    return p


def add_rich_para(doc, segments, alignment=None, space_after=Pt(6)):
    """segments: list of (text, font_size, bold, color)"""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = Pt(20)
    for text, font_size, bold, color in segments:
        add_run(p, text, font_size=font_size, bold=bold, color=color)
    return p


def add_table(doc, headers, rows, col_alignments=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, header, font_size=Pt(10), bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(cell, YIWU_RED_HEX)

    # 数据行
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            align = WD_ALIGN_PARAGRAPH.CENTER
            if col_alignments and c_idx < len(col_alignments):
                align = col_alignments[c_idx]
            p.alignment = align
            is_bold = "合计" in str(cell_text) or "总计" in str(cell_text) or "月度总计" in str(cell_text)
            add_run(p, str(cell_text), font_size=Pt(9.5), bold=is_bold)
            if r_idx % 2 == 1:
                set_cell_shading(cell, LIGHT_RED_HEX)

    table.autofit = True
    doc.add_paragraph()
    return table


def add_page_header(section, text):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, text, font_size=Pt(8), color=RGBColor(0x99, 0x99, 0x99))
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


def add_page_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)
    add_run(p, "第 ", font_size=Pt(8), color=RGBColor(0x99, 0x99, 0x99))
    # 页码字段
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run1 = p.add_run()
    run1._element.append(fldChar1)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2 = p.add_run()
    run2._element.append(instrText)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3 = p.add_run()
    run3._element.append(fldChar2)
    add_run(p, " 页", font_size=Pt(8), color=RGBColor(0x99, 0x99, 0x99))


# ── 封面页 ──────────────────────────────────────────

def create_cover_page(doc):
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "义乌小商品出海智能体-OPC", font_name=FONT_HEADING, font_size=Pt(32), bold=True,
            color=RGBColor(0xD4, 0x27, 0x2C))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    add_run(p, "OPC数字商业运营组 参赛商业计划书", font_name=FONT_HEADING, font_size=Pt(22), bold=True,
            color=RGBColor(0x33, 0x33, 0x33))

    # 分隔线
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(20)
    add_run(p, "━" * 30, font_size=Pt(12), color=RGBColor(0xD4, 0x27, 0x2C))

    info_lines = [
        "V2.0 冠军版",
        "参赛赛道：2026\"直通乌镇\"全球互联网大赛 OPC特色赛",
        "日期：2026年6月",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        add_run(p, line, font_size=Pt(12), color=RGBColor(0x55, 0x55, 0x55))

    doc.add_page_break()


# ── 目录页 ──────────────────────────────────────────

def create_toc_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    add_run(p, "目  录", font_name=FONT_HEADING, font_size=Pt(22), bold=True,
            color=RGBColor(0xD4, 0x27, 0x2C))

    toc_items = [
        ("第一章  项目概述", 1),
        ("1.1 项目简介", 2),
        ("1.2 OPC模式定义", 2),
        ("第二章  AI提效可视化对比（重点）", 1),
        ("2.1 传统模式 vs OPC模式人力对比表", 2),
        ("2.2 时间成本对比图", 2),
        ("2.3 费用对比表", 2),
        ("第三章  AI提效深度与广度（重点）", 1),
        ("3.1 单个AI工具的使用深度", 2),
        ("3.2 多Agent协调能力（广度）", 2),
        ("第四章  产品矩阵", 1),
        ("4.1 七大AI Agent产品矩阵", 2),
        ("4.2 四档定价矩阵", 2),
        ("4.3 三级增长飞轮", 2),
        ("第五章  义乌独有壁垒", 1),
        ("第六章  商业模式与财务", 1),
        ("第七章  团队与愿景", 1),
    ]

    for title, level in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = Pt(22)
        if level == 1:
            p.paragraph_format.left_indent = Cm(0)
            add_run(p, title, font_size=Pt(12), bold=True, color=RGBColor(0xD4, 0x27, 0x2C))
        else:
            p.paragraph_format.left_indent = Cm(1.2)
            add_run(p, title, font_size=Pt(10.5), color=RGBColor(0x44, 0x44, 0x44))

    doc.add_page_break()


# ── 第一章：项目概述 ──────────────────────────────────

def create_chapter1(doc):
    add_heading(doc, "第一章  项目概述", level=1)

    add_heading(doc, "1.1 项目简介", level=2)
    add_para(doc,
        "义乌小商品出海智能体-OPC，是基于义乌7.5万商户、210万SKU、39城政策经验的AI全链路出海平台。"
        "通过7个专业AI Agent协同工作，实现从市场洞察、智能选品、供应链匹配、跨境内容生成、合规审查、"
        "智能客服到政策复制的全流程自动化，让1个人即可完成传统10人团队的出海工作，成本降低90%以上。")

    add_heading(doc, "1.2 OPC模式定义", level=2)
    add_para(doc, "OPC = One Person Company，即\"一人公司\"模式。", bold=True)
    add_para(doc,
        "核心公式：1人 + 7个AI Agent = 传统10人团队")
    add_para(doc,
        "OPC模式不是简单的工具堆砌，而是通过LangGraph状态机将7个AI Agent编排为全链路工作流，"
        "前一步的输出自动成为下一步的输入，实现真正的\"一人即团队\"。传统模式下，出海需要市场、选品、"
        "供应链、内容、合规、客服、政策等7个环节14人持续协作；OPC模式下，1个人+7个AI Agent即可"
        "完成全部工作，效率提升10倍以上，成本降低90%以上。")

    doc.add_page_break()


# ── 第二章：AI提效可视化对比 ──────────────────────────

def create_chapter2(doc):
    add_heading(doc, "第二章  AI提效可视化对比", level=1)
    add_para(doc, "本章是OPC赛道的核心展示，通过三组可视化对比，直观呈现OPC模式的提效效果。",
             bold=True, color=RGBColor(0xD4, 0x27, 0x2C))

    # 2.1 人力对比表
    add_heading(doc, "2.1 传统模式 vs OPC模式人力对比表", level=2)
    add_para(doc, "以下对比展示了传统出海模式与OPC模式在各环节的人力投入差异：")

    headers = ["环节", "传统模式", "OPC模式", "提效"]
    rows = [
        ["市场调研", "2人×2周", "1人×5分钟", "99.9%"],
        ["选品分析", "2人×1周", "1人×3分钟", "99.8%"],
        ["供应链对接", "3人×2周", "1人×10分钟", "99.5%"],
        ["内容创作", "2人×1周", "1人×5分钟", "99.8%"],
        ["合规审查", "1人×3天", "1人×2分钟", "99.7%"],
        ["客户服务", "3人×持续", "1人×0分钟(自动)", "100%"],
        ["政策研究", "1人×1周", "1人×3分钟", "99.7%"],
        ["合计", "14人×持续", "1人+7Agent", "成本降90%"],
    ]
    add_table(doc, headers, rows, col_alignments=[
        WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER
    ])

    # 2.2 时间成本对比图
    add_heading(doc, "2.2 时间成本对比图", level=2)
    add_para(doc, "用表格模拟柱状图，直观对比传统模式与OPC模式的全流程时间消耗：")

    # 模拟柱状图 - 传统模式
    add_para(doc, "▎传统模式出海全流程：约8周", bold=True, color=RGBColor(0x99, 0x99, 0x99))
    bar_headers = ["阶段", "耗时", "可视化"]
    bar_rows = [
        ["市场调研", "2周", "████████████████████"],
        ["选品分析", "1周", "██████████"],
        ["供应链对接", "2周", "████████████████████"],
        ["内容创作", "1周", "██████████"],
        ["合规审查", "3天", "█████"],
        ["客户服务", "持续", "████████████████████████████"],
        ["政策研究", "1周", "██████████"],
    ]
    add_table(doc, bar_headers, bar_rows, col_alignments=[
        WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT
    ])

    # 模拟柱状图 - OPC模式
    add_para(doc, "▎OPC模式出海全流程：约1天", bold=True, color=RGBColor(0xD4, 0x27, 0x2C))
    opc_bar_headers = ["阶段", "耗时", "可视化"]
    opc_bar_rows = [
        ["市场调研", "5分钟", "▌"],
        ["选品分析", "3分钟", "▌"],
        ["供应链对接", "10分钟", "▌"],
        ["内容创作", "5分钟", "▌"],
        ["合规审查", "2分钟", "▌"],
        ["客户服务", "0分钟(自动)", "▌"],
        ["政策研究", "3分钟", "▌"],
    ]
    add_table(doc, opc_bar_headers, opc_bar_rows, col_alignments=[
        WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT
    ])

    # 对比总结
    add_rich_para(doc, [
        ("对比结论：传统模式 8周 → OPC模式 1天，", Pt(12), True, RGBColor(0xD4, 0x27, 0x2C)),
        ("时间缩短 98.2%", Pt(12), True, RGBColor(0xD4, 0x27, 0x2C)),
    ], alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    # 2.3 费用对比表
    add_heading(doc, "2.3 费用对比表", level=2)
    add_para(doc, "月度运营费用对比，OPC模式在人力、工具、合规、翻译等方面全面降本：")

    fee_headers = ["项目", "传统模式（月）", "OPC模式（月）", "节省"]
    fee_rows = [
        ["人力成本", "14人×8000=112,000元", "1人×8000=8,000元", "93%"],
        ["工具订阅", "分散工具约5,000元", "OPC平台999元", "80%"],
        ["合规咨询", "约10,000元", "含在平台内", "100%"],
        ["翻译服务", "约8,000元", "含在平台内", "100%"],
        ["月度总计", "135,000元", "8,999元", "93%"],
    ]
    add_table(doc, fee_headers, fee_rows, col_alignments=[
        WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER
    ])

    # 费用对比视觉强化
    add_rich_para(doc, [
        ("💰 月度节省：", Pt(14), True, RGBColor(0x33, 0x33, 0x33)),
        ("135,000元 → 8,999元", Pt(14), True, RGBColor(0xD4, 0x27, 0x2C)),
        ("，年节省超150万元", Pt(14), True, RGBColor(0x33, 0x33, 0x33)),
    ], alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    doc.add_page_break()


# ── 第三章：AI提效深度与广度 ──────────────────────────

def create_chapter3(doc):
    add_heading(doc, "第三章  AI提效深度与广度", level=1)
    add_para(doc, "OPC赛道不仅关注AI工具的数量，更关注每个AI工具的使用深度和多工具之间的协调广度。",
             bold=True, color=RGBColor(0xD4, 0x27, 0x2C))

    # 3.1 单个AI工具的使用深度
    add_heading(doc, "3.1 单个AI工具的使用深度", level=2)
    add_para(doc, "7个AI Agent各自拥有深度专业能力，不是简单的\"套壳ChatGPT\"，而是基于行业数据、"
             "专业知识图谱和领域模型的深度AI应用：")

    agent_headers = ["Agent", "核心能力", "单独提效", "数据源", "技术"]
    agent_rows = [
        ["市场洞察", "全球市场趋势分析", "提效50倍", "Amazon/Shopee/义乌指数", "LangGraph+RAG"],
        ["智能选品", "210万SKU智能匹配", "提效100倍", "1688/义乌小商品城", "多维度评分模型"],
        ["供应链匹配", "7.5万商铺精准对接", "提效200倍", "义乌商铺数据库", "向量检索+匹配"],
        ["跨境内容", "8语言4平台内容生成", "提效80倍", "LLM+平台规则", "DashScope Qwen"],
        ["合规助手", "5国合规+1039+RCEP", "提效300倍", "法规数据库", "知识图谱+LLM"],
        ["智能客服", "7×24多语言客服", "提效∞(无人)", "FAQ+产品库", "情绪识别+LLM"],
        ["政策复制", "39城政策解读复制", "提效100倍", "39城政策库", "政策图谱+LLM"],
    ]
    add_table(doc, agent_headers, agent_rows, col_alignments=[
        WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT
    ])

    # 深度解读
    add_para(doc, "深度解读：", bold=True)
    depth_items = [
        "市场洞察Agent：不是简单的关键词搜索，而是基于RAG架构，融合Amazon、Shopee实时数据和义乌指数历史数据，输出结构化市场趋势报告",
        "智能选品Agent：基于210万SKU的多维度评分模型，综合考虑市场需求、利润率、供应链匹配度、合规风险等8个维度",
        "供应链匹配Agent：7.5万商铺的向量检索+精准匹配，支持按品类、价格、MOQ、交期等多条件筛选",
        "跨境内容Agent：基于DashScope Qwen大模型，支持8种语言、4大电商平台的内容风格适配，自动生成标题、描述、关键词",
        "合规助手Agent：基于5国法规知识图谱+1039市场采购贸易模式+RCEP规则，实现合规风险实时预警",
        "智能客服Agent：7×24小时多语言自动回复，支持情绪识别和复杂问题升级，真正实现\"无人客服\"",
        "政策复制Agent：39城政策图谱+红利自动计算，将义乌经验一键复制到全国",
    ]
    for item in depth_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = Pt(18)
        p.paragraph_format.left_indent = Cm(0.8)
        add_run(p, "• " + item, font_size=Pt(9.5))

    # 3.2 多Agent协调能力
    add_heading(doc, "3.2 多Agent协调能力（广度）", level=2)
    add_para(doc, "7个Agent不是独立工具，而是通过LangGraph状态机协调的全链路工作流，前一步的输出自动成为下一步的输入。")

    add_para(doc, "全链路7步工作流：", bold=True)
    workflow_steps = [
        "1. 市场洞察 → 发现目标市场机会",
        "2. 智能选品 → 匹配最优商品",
        "3. 供应链匹配 → 对接义乌商铺",
        "4. 跨境内容 → 生成8语言内容",
        "5. 合规审查 → 确保合规出海",
        "6. 智能客服 → 自动客户服务",
        "7. 政策复制 → 39城经验推广",
    ]
    for step in workflow_steps:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = Pt(18)
        p.paragraph_format.left_indent = Cm(1.0)
        add_run(p, step, font_size=Pt(10.5))

    add_para(doc, "")

    # 协调效果对比表
    add_para(doc, "协调效果对比：", bold=True)
    coord_headers = ["协调场景", "无协调（分别使用）", "有协调（OPC全链路）", "额外提效"]
    coord_rows = [
        ["选品→供应链", "手动复制粘贴", "自动传递选品结果", "3倍"],
        ["内容→合规", "先生成再审查", "生成时实时合规检查", "2倍"],
        ["全链路", "7个工具分别操作", "一键全链路执行", "10倍"],
    ]
    add_table(doc, coord_headers, coord_rows, col_alignments=[
        WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER
    ])

    add_rich_para(doc, [
        ("关键结论：7个Agent协调使用，比分别使用额外提效", Pt(11), True, RGBColor(0x33, 0x33, 0x33)),
        ("10倍", Pt(11), True, RGBColor(0xD4, 0x27, 0x2C)),
        ("，这就是OPC模式的核心竞争力", Pt(11), True, RGBColor(0x33, 0x33, 0x33)),
    ], alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    doc.add_page_break()


# ── 第四章：产品矩阵 ──────────────────────────────────

def create_chapter4(doc):
    add_heading(doc, "第四章  产品矩阵", level=1)

    # 4.1 七大AI Agent产品矩阵
    add_heading(doc, "4.1 七大AI Agent产品矩阵", level=2)

    matrix_headers = ["Agent", "输入", "输出", "适用场景"]
    matrix_rows = [
        ["市场洞察Agent", "目标品类", "市场趋势报告", "市场选择"],
        ["智能选品Agent", "市场需求", "选品推荐+评分", "商品选择"],
        ["供应链匹配Agent", "选品结果", "商铺匹配+报价", "采购对接"],
        ["跨境内容Agent", "商品信息", "8语言4平台内容", "营销推广"],
        ["合规助手Agent", "目标国家", "合规清单+1039方案", "合规出海"],
        ["智能客服Agent", "客户咨询", "多语言自动回复", "客户服务"],
        ["政策复制Agent", "城市信息", "政策解读+红利计算", "经验复制"],
    ]
    add_table(doc, matrix_headers, matrix_rows, col_alignments=[
        WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT
    ])

    # 4.2 四档定价矩阵
    add_heading(doc, "4.2 四档定价矩阵", level=2)

    pricing_headers = ["版本", "月费", "Agent数量", "目标客户"]
    pricing_rows = [
        ["试水版", "199元", "3个", "个体商户"],
        ["起航版", "499元", "5个", "小微企业"],
        ["远航版", "999元", "7个", "成长企业"],
        ["领航版", "2999元", "7个+专属", "大型企业"],
    ]
    add_table(doc, pricing_headers, pricing_rows, col_alignments=[
        WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER
    ])

    # 4.3 三级增长飞轮
    add_heading(doc, "4.3 三级增长飞轮", level=2)

    flywheel_headers = ["级别", "覆盖范围", "规模", "预期收入"]
    flywheel_rows = [
        ["第一级", "义乌本地7.5万商户", "7.5万", "月MRR 144.7万"],
        ["第二级", "全国39城210万企业", "210万", "24个月净收入3,962万"],
        ["第三级", "全国县域5000万企业", "5000万", "长期规模化增长"],
    ]
    add_table(doc, flywheel_headers, flywheel_rows, col_alignments=[
        WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER
    ])

    doc.add_page_break()


# ── 第五章：义乌独有壁垒 ──────────────────────────────

def create_chapter5(doc):
    add_heading(doc, "第五章  义乌独有壁垒", level=1)

    add_para(doc, "三大独有壁垒 + 义乌发展经验国家级制度性壁垒，构成不可复制的竞争护城河：")

    barrier_headers = ["壁垒类型", "具体内容", "竞争对手可复制性"]
    barrier_rows = [
        ["数据壁垒", "7.5万商铺真实数据+210万SKU+义乌指数20年历史数据", "极难复制"],
        ["政策壁垒", "义乌综试区+1039市场采购贸易模式+RCEP先行区", "国家级制度性壁垒"],
        ["经验壁垒", "义乌\"六个坚持\"发展经验+39城政策复制经验", "需10年+积累"],
        ["生态壁垒", "义新欧班列+海外仓+跨境电商产业园完整生态", "需政府级资源"],
    ]
    add_table(doc, barrier_headers, barrier_rows, col_alignments=[
        WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER
    ])

    add_para(doc, "")
    add_para(doc, "义乌发展经验——国家级制度性壁垒：", bold=True, color=RGBColor(0xD4, 0x27, 0x2C))
    add_para(doc,
        "义乌\"六个坚持\"发展经验（坚持兴商建市、坚持产业联动、坚持城乡统筹、坚持和谐发展、"
        "坚持共创共富、坚持党建引领）已被提升为国家级制度性经验，这是任何竞争对手无法复制的"
        "根本性壁垒。我们的政策复制Agent正是基于这一独特优势，将义乌经验系统化、AI化，"
        "实现39城一键复制。")

    doc.add_page_break()


# ── 第六章：商业模式与财务 ──────────────────────────────

def create_chapter6(doc):
    add_heading(doc, "第六章  商业模式与财务", level=1)

    add_heading(doc, "6.1 收入模式", level=2)
    add_para(doc, "三大收入引擎：")
    revenue_items = [
        "SaaS订阅：试水版199元/月起，覆盖7.5万商户到5000万企业",
        "政府采购：义乌综试区+39城政策复制，单城年费50-200万",
        "城市合伙人：39城独家运营权，分成模式快速扩张",
    ]
    for item in revenue_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = Pt(18)
        p.paragraph_format.left_indent = Cm(0.8)
        add_run(p, "• " + item, font_size=Pt(10.5))

    add_heading(doc, "6.2 财务预测", level=2)

    finance_headers = ["指标", "12个月", "24个月", "36个月"]
    finance_rows = [
        ["月MRR", "144.7万", "580万", "1,500万"],
        ["累计收入", "868万", "3,962万", "1.2亿"],
        ["付费用户", "1,447", "5,800", "15,000"],
        ["城市覆盖", "义乌+5城", "39城", "100城"],
        ["毛利率", "78%", "82%", "85%"],
    ]
    add_table(doc, finance_headers, finance_rows, col_alignments=[
        WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER
    ])

    add_rich_para(doc, [
        ("核心指标：12个月MRR ", Pt(12), True, RGBColor(0x33, 0x33, 0x33)),
        ("144.7万", Pt(12), True, RGBColor(0xD4, 0x27, 0x2C)),
        ("，39城24个月净收入 ", Pt(12), True, RGBColor(0x33, 0x33, 0x33)),
        ("3,962万", Pt(12), True, RGBColor(0xD4, 0x27, 0x2C)),
    ], alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    doc.add_page_break()


# ── 第七章：团队与愿景 ──────────────────────────────

def create_chapter7(doc):
    add_heading(doc, "第七章  团队与愿景", level=1)

    add_heading(doc, "7.1 OPC模式团队", level=2)
    add_para(doc, "团队核心：1人 + 7个AI Agent", bold=True, color=RGBColor(0xD4, 0x27, 0x2C))

    team_headers = ["角色", "传统团队", "OPC模式", "说明"]
    team_rows = [
        ["市场分析师", "2人", "市场洞察Agent", "RAG+实时数据分析"],
        ["选品专家", "2人", "智能选品Agent", "210万SKU多维评分"],
        ["供应链经理", "3人", "供应链匹配Agent", "7.5万商铺精准匹配"],
        ["内容运营", "2人", "跨境内容Agent", "8语言4平台自动生成"],
        ["合规顾问", "1人", "合规助手Agent", "5国法规+1039+RCEP"],
        ["客服团队", "3人", "智能客服Agent", "7×24多语言无人客服"],
        ["政策研究员", "1人", "政策复制Agent", "39城政策图谱+红利计算"],
    ]
    add_table(doc, team_headers, team_rows, col_alignments=[
        WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT
    ])

    add_heading(doc, "7.2 创始人", level=2)
    add_para(doc,
        "冯亦根——浙江省首批产业教授、高级工程师、工信部认证首席数据官，深耕义乌小商品产业20年，"
        "深度理解商户痛点和政策体系，是义乌数字化转型的核心推动者。")

    add_heading(doc, "7.3 愿景", level=2)
    add_rich_para(doc, [
        ("让小商品照亮全球", Pt(18), True, RGBColor(0xD4, 0x27, 0x2C)),
    ], alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    add_para(doc,
        "我们相信，OPC模式不仅是义乌的答案，更是中国5000万县域企业出海的答案。"
        "通过AI赋能，让每一个小商户都拥有媲美大企业的出海能力，"
        "让义乌的小商品通过智能化的方式照亮全球每一个角落。")


# ── 主函数 ──────────────────────────────────────────

def main():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_BODY
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)

    # 设置页面边距
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    # 页眉页脚
    add_page_header(section, "义乌小商品出海智能体-OPC  OPC赛道商业计划书")
    add_page_footer(section)

    # 封面
    create_cover_page(doc)

    # 目录
    create_toc_page(doc)

    # 正文
    create_chapter1(doc)
    create_chapter2(doc)
    create_chapter3(doc)
    create_chapter4(doc)
    create_chapter5(doc)
    create_chapter6(doc)
    create_chapter7(doc)

    # 保存
    doc.save(str(DOCX_PATH))
    print(f"Word文档已成功生成：{DOCX_PATH}")
    print(f"文件大小：{DOCX_PATH.stat().st_size / 1024:.1f} KB")


if __name__ == '__main__':
    main()
